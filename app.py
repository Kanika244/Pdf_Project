import streamlit as st
import io
import os
import zipfile
import tempfile
from pathlib import Path
import pdfplumber
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pypdf import PdfReader, PdfWriter
from reportlab.pdfgen import canvas as rl_canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.colors import HexColor
from pdf2image import convert_from_bytes
import pytesseract

# ─── Page Config ────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="UniPDF – Free PDF Tools for Students",
    layout="wide",
    initial_sidebar_state="expanded",
)


# ─── Sidebar Navigation ──────────────────────────────────────────────────────
with st.sidebar:
    st.markdown("""
    <div style='text-align:center; padding: 1rem 0 1.5rem'>
        <div style='font-size:2.5rem'>📄</div>
        <div style='font-size:1.3rem; font-weight:700; color:white'>UniPDF</div>
        <div style='font-size:0.78rem; opacity:0.7; color:#94a3b8'>Free PDF Tools for Students</div>
    </div>
    """, unsafe_allow_html=True)

    tool = st.radio("", [
        "Home",
        "Merge PDFs",
        "Split PDF",
        "Compress PDF",
        "Rotate Pages",
        "PDF → Word",
        "Add Watermark",
        "Add Page Numbers",
        "Extract Text",
        "Extract Tables",
        "Reorder Pages",
        "OCR (Scan to Text)",
        "Named Entity Recognition",
    ], label_visibility="collapsed")

    st.markdown("---")
    st.markdown("""
    <div style='font-size:0.78rem; color:#94a3b8; text-align:center'>
         100% Free &nbsp;|&nbsp; No data stored<br>
        Built for university students
    </div>
    """, unsafe_allow_html=True)


# ─── Helper: download button ─────────────────────────────────────────────────
def download_btn(data: bytes, filename: str, label: str = "Download Result"):
    mime = "application/pdf" if filename.endswith(".pdf") else "application/octet-stream"
    if filename.endswith(".docx"):
        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if filename.endswith(".zip"):
        mime = "application/zip"
    if filename.endswith(".txt"):
        mime = "text/plain"
    if filename.endswith(".csv"):
        mime = "text/csv"
    st.download_button(label, data=data, file_name=filename, mime=mime, use_container_width=True)


def section_header(title, subtitle, premium=False):
    label = title + (" — Premium (Free for Students)" if premium else "")
    st.title(label)
    st.caption(subtitle)
    st.divider()


if tool == "Home":
    st.title("UniPDF")
    st.subheader("The all-in-one PDF toolkit built exclusively for university students.")
    st.write("Every premium feature — completely free. No sign-up required.")
    st.divider()

    col1, col2, col3 = st.columns(3)
    with col1:
        st.subheader("Merge PDFs")
        st.write("Combine multiple PDFs into one document.")
        st.subheader("Split PDF")
        st.write("Extract pages or split into chunks.")
        st.subheader("Compress PDF")
        st.write("Reduce file size significantly.")
    with col2:
        st.subheader("PDF → Word")
        st.write("Convert PDF to an editable DOCX file.")
        st.subheader("PDF ↔ Images")
        st.write("Convert between PDF and images.")
        st.subheader("Watermark")
        st.write("Add text watermarks to any PDF.")
    with col3:
        st.subheader("Protect / Unlock")
        st.write("Password protect or unlock PDFs.")
        st.subheader("OCR")
        st.write("Extract text from scanned PDFs.")
        st.subheader("Extract Tables")
        st.write("Pull tables to CSV or Excel.")

    st.divider()
    st.info("Pick a tool from the sidebar to get started. All tools are free, no sign-up required.")


elif tool == "Merge PDFs":
    section_header("Merge PDFs", "Combine multiple PDF files into a single document.")

    files = st.file_uploader("Upload PDFs to merge", type="pdf", accept_multiple_files=True)

    if files:
        st.markdown(f"**{len(files)} file(s) uploaded** (will be merged in order shown)")
        for i, f in enumerate(files):
            st.markdown(f"&nbsp;&nbsp;📄 {i+1}. {f.name}")

        if st.button("Merge PDFs", use_container_width=True):
            writer = PdfWriter()
            try:
                for f in files:
                    reader = PdfReader(io.BytesIO(f.read()))
                    for page in reader.pages:
                        writer.add_page(page)
                buf = io.BytesIO()
                writer.write(buf)
                st.success(f"Merged {len(files)} PDFs successfully!")
                download_btn(buf.getvalue(), "merged.pdf", "Download Merged PDF")
            except Exception as e:
                st.error(f"Error: {e}")
    else:
        st.info("Upload 2 or more PDF files to merge them.")



elif tool == "Split PDF":
    section_header("Split PDF", "Extract specific pages or split into individual pages / equal chunks.")

    file = st.file_uploader("Upload a PDF to split", type="pdf")

    if file:
        reader = PdfReader(io.BytesIO(file.read()))
        total = len(reader.pages)
        st.info(f"**{file.name}** — {total} pages")

        split_mode = st.radio("Split mode", [
            "Extract specific pages",
            "Split every N pages",
            "Split into individual pages"
        ])

        if split_mode == "Extract specific pages":
            page_input = st.text_input("Enter page numbers / ranges (e.g. 1,3,5-8,10)", placeholder="1,3,5-8,10")
            if st.button("Extract Pages", use_container_width=True):
                def parse_pages(s, total):
                    pages = set()
                    for part in s.split(","):
                        part = part.strip()
                        if "-" in part:
                            a, b = part.split("-")
                            pages.update(range(int(a)-1, min(int(b), total)))
                        else:
                            p = int(part) - 1
                            if 0 <= p < total:
                                pages.add(p)
                    return sorted(pages)
                try:
                    idxs = parse_pages(page_input, total)
                    writer = PdfWriter()
                    for i in idxs:
                        writer.add_page(reader.pages[i])
                    buf = io.BytesIO()
                    writer.write(buf)
                    st.success(f"Extracted {len(idxs)} pages!")
                    download_btn(buf.getvalue(), "extracted_pages.pdf")
                except Exception as e:
                    st.error(f"Error: {e}")

        elif split_mode == "Split every N pages":
            n = st.number_input("Pages per chunk", min_value=1, max_value=total, value=min(5, total))
            if st.button("Split", use_container_width=True):
                zbuf = io.BytesIO()
                with zipfile.ZipFile(zbuf, "w") as zf:
                    chunks = list(range(0, total, int(n)))
                    for idx, start in enumerate(chunks):
                        end = min(start + int(n), total)
                        writer = PdfWriter()
                        for i in range(start, end):
                            writer.add_page(reader.pages[i])
                        pbuf = io.BytesIO()
                        writer.write(pbuf)
                        zf.writestr(f"part_{idx+1}_pages_{start+1}-{end}.pdf", pbuf.getvalue())
                st.success(f"Split into {len(chunks)} parts!")
                download_btn(zbuf.getvalue(), "split_parts.zip", "Download All Parts (ZIP)")

        else:
            if st.button("Split into Individual Pages", use_container_width=True):
                zbuf = io.BytesIO()
                with zipfile.ZipFile(zbuf, "w") as zf:
                    for i, page in enumerate(reader.pages):
                        writer = PdfWriter()
                        writer.add_page(page)
                        pbuf = io.BytesIO()
                        writer.write(pbuf)
                        zf.writestr(f"page_{i+1:03d}.pdf", pbuf.getvalue())
                st.success(f"Split into {total} individual pages!")
                download_btn(zbuf.getvalue(), "individual_pages.zip", "Download All Pages (ZIP)")
    else:
        st.info("Upload a PDF file to split it.")


elif tool == "Compress PDF":
    section_header("Compress PDF", "Reduce the file size of your PDFs.", premium=True)

    file = st.file_uploader("Upload a PDF to compress", type="pdf")

    if file:
        original_size = len(file.getvalue())
        st.info(f"Original size: **{original_size / 1024:.1f} KB**")

        level = st.select_slider(
            "Compression level",
            options=["Low (best quality)", "Medium (balanced)", "High (smallest size)"],
            value="Medium (balanced)"
        )

        if st.button("Compress PDF", use_container_width=True):
            try:
                reader = PdfReader(io.BytesIO(file.getvalue()))
                writer = PdfWriter()
                for page in reader.pages:
                    writer.add_page(page)


                if level != "Low (best quality)":
                        for page in writer.pages:
                            page.compress_content_streams()
                   
                writer.add_metadata(reader.metadata or {})
                buf = io.BytesIO()
                writer.write(buf)
                compressed = buf.getvalue()
                new_size = len(compressed)
                savings = (1 - new_size / original_size) * 100

                col1, col2, col3 = st.columns(3)
                col1.metric("Original", f"{original_size/1024:.1f} KB")
                col2.metric("Compressed", f"{new_size/1024:.1f} KB")
                col3.metric("Savings", f"{savings:.1f}%", delta=f"-{savings:.1f}%")

                st.success("Compression complete!")
                download_btn(compressed, "compressed.pdf", "Download Compressed PDF")
            except Exception as e:
                st.error(f"Error: {e}")
    else:
        st.info("Upload a PDF to compress it.")



elif tool == "Rotate Pages":
    section_header("Rotate Pages", "Rotate specific pages or all pages in a PDF.")

    file = st.file_uploader("Upload a PDF", type="pdf")

    if file:
        reader = PdfReader(io.BytesIO(file.read()))
        total = len(reader.pages)
        st.info(f"**{file.name}** — {total} pages")

        col1, col2 = st.columns(2)
        with col1:
            rotate_scope = st.radio("Apply rotation to", ["All pages", "Specific pages"])
        with col2:
            angle = st.selectbox("Rotation angle", [90, 180, 270], index=0)

        pages_input = ""
        if rotate_scope == "Specific pages":
            pages_input = st.text_input("Page numbers / ranges (e.g. 1,3,5-8)", placeholder="1,3,5-8")

        if st.button("Rotate", use_container_width=True):
            def parse_pages(s, total):
                pages = set()
                for part in s.split(","):
                    part = part.strip()
                    if "-" in part:
                        a, b = part.split("-")
                        pages.update(range(int(a)-1, min(int(b), total)))
                    elif part:
                        p = int(part) - 1
                        if 0 <= p < total:
                            pages.add(p)
                return pages

            try:
                rotate_set = set(range(total)) if rotate_scope == "All pages" else parse_pages(pages_input, total)
                writer = PdfWriter()
                for i, page in enumerate(reader.pages):
                    writer.add_page(page)
                for i, page in enumerate(writer.pages):
                    if i in rotate_set:
                        page.rotate(angle)
                buf = io.BytesIO()
                writer.write(buf)
                st.success(f"Rotated {len(rotate_set)} pages by {angle}°!")
                download_btn(buf.getvalue(), "rotated.pdf", "Download Rotated PDF")
            except Exception as e:
                st.error(f"Error: {e}")
    else:
        st.info("Upload a PDF to rotate its pages.")


elif tool == "PDF → Word":
    section_header("PDF → Word", "Convert a PDF to an editable Microsoft Word (.docx) document.", premium=True)

    file = st.file_uploader("Upload a PDF", type="pdf")

    if file:
        if st.button("Convert to Word", use_container_width=True):
            try:
                doc = Document()
                title = doc.add_heading(Path(file.name).stem, 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER

                with pdfplumber.open(io.BytesIO(file.read())) as pdf:
                    for i, page in enumerate(pdf.pages):
                        if i > 0:
                            doc.add_page_break()
                        doc.add_heading(f"Page {i+1}", level=2)
                        text = page.extract_text()
                        if text:
                            for para_text in text.split("\n"):
                                if para_text.strip():
                                    doc.add_paragraph(para_text.strip())
                        for tbl_data in page.extract_tables():
                            if tbl_data:
                                doc.add_paragraph("Table:", style='Intense Quote')
                                rows = [r for r in tbl_data if r]
                                if rows:
                                    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
                                    t.style = "Table Grid"
                                    for r_idx, row in enumerate(rows):
                                        for c_idx, cell in enumerate(row):
                                            t.cell(r_idx, c_idx).text = str(cell or "")

                buf = io.BytesIO()
                doc.save(buf)
                st.success("Converted to Word document!")
                download_btn(buf.getvalue(), f"{Path(file.name).stem}.docx", "Download Word Document")
            except Exception as e:
                st.error(f"Error: {e}")
    else:
        st.info("Upload a PDF to convert it to Word format.")


elif tool == "Add Watermark":
    section_header("Add Watermark", "Stamp a text watermark on every page of your PDF.", premium=True)

    file = st.file_uploader("Upload a PDF", type="pdf")

    if file:
        col1, col2 = st.columns(2)
        wm_text    = col1.text_input("Watermark text", value="CONFIDENTIAL")
        wm_angle   = col2.slider("Angle (degrees)", 0, 90, 45)
        col3, col4 = st.columns(2)
        wm_opacity = col3.slider("Opacity", 0.05, 0.5, 0.15, step=0.05)
        wm_color   = col4.color_picker("Color", "#FF0000")

        if st.button("Add Watermark", use_container_width=True):
            try:
                # Build watermark page
                wm_buf = io.BytesIO()
                c = rl_canvas.Canvas(wm_buf, pagesize=letter)
                w, h = letter
                c.saveState()
                c.setFillColor(HexColor(wm_color))
                c.setFillAlpha(wm_opacity)
                c.setFont("Helvetica-Bold", 60)
                c.translate(w / 2, h / 2)
                c.rotate(wm_angle)
                c.drawCentredString(0, 0, wm_text)
                c.restoreState()
                c.save()

                wm_page = PdfReader(wm_buf).pages[0]

                reader = PdfReader(io.BytesIO(file.read()))
                writer = PdfWriter()

             
                for page in reader.pages:
                    writer.add_page(page)

               
                for page in writer.pages:
                    page.merge_page(wm_page)

                buf = io.BytesIO()
                writer.write(buf)
                st.success("Watermark added to all pages!")
                download_btn(buf.getvalue(), "watermarked.pdf", "Download Watermarked PDF")
            except Exception as e:
                st.error(f"Error: {e}")
    else:
        st.info("Upload a PDF to add a watermark.")


elif tool == "Add Page Numbers":
    section_header("Add Page Numbers", "Automatically add page numbers to your PDF.", premium=True)

    file = st.file_uploader("Upload a PDF", type="pdf")

    if file:
        col1, col2, col3 = st.columns(3)
        position  = col1.selectbox("Position", ["Bottom Center", "Bottom Right", "Bottom Left", "Top Center"])
        start_num = col2.number_input("Start from page number", min_value=1, value=1)
        font_size = col3.slider("Font size", 8, 18, 11)

        if st.button("Add Page Numbers", use_container_width=True):
            try:
                reader = PdfReader(io.BytesIO(file.read()))
                writer = PdfWriter()

                for page in reader.pages:
                    writer.add_page(page)

                
                for i, page in enumerate(writer.pages):
                    pw = float(page.mediabox.width)
                    ph = float(page.mediabox.height)

                    num_buf = io.BytesIO()
                    c = rl_canvas.Canvas(num_buf, pagesize=(pw, ph))
                    c.setFont("Helvetica", font_size)
                    c.setFillColorRGB(0.3, 0.3, 0.3)
                    pg_label = str(i + int(start_num))
                    margin = 30
                    if position == "Bottom Center":
                        c.drawCentredString(pw / 2, margin, pg_label)
                    elif position == "Bottom Right":
                        c.drawRightString(pw - margin, margin, pg_label)
                    elif position == "Bottom Left":
                        c.drawString(margin, margin, pg_label)
                    else:  # Top Center
                        c.drawCentredString(pw / 2, ph - margin, pg_label)
                    c.save()

                    num_page = PdfReader(num_buf).pages[0]
                    page.merge_page(num_page)

                buf = io.BytesIO()
                writer.write(buf)
                st.success(f"Page numbers added to {len(reader.pages)} pages!")
                download_btn(buf.getvalue(), "numbered.pdf", "Download Numbered PDF")
            except Exception as e:
                st.error(f"Error: {e}")
    else:
        st.info("Upload a PDF to add page numbers.")


elif tool == "Extract Text":
    section_header("Extract Text", "Extract all text content from a PDF document.")

    file = st.file_uploader("Upload a PDF", type="pdf")

    if file:
        col1, col2 = st.columns(2)
        method     = col1.radio("Extraction method", ["pdfplumber (layout-aware)", "pypdf (fast)"])
        page_range = col2.radio("Pages", ["All pages", "Specific pages"])

        page_input = ""
        if page_range == "Specific pages":
            page_input = st.text_input("Page numbers / ranges (e.g. 1-5, 8)", placeholder="1-5, 8")

        if st.button("Extract Text", use_container_width=True):
            try:
                text_pages = []

                if "pdfplumber" in method:
                    with pdfplumber.open(io.BytesIO(file.read())) as pdf:
                        total = len(pdf.pages)
                        idxs = range(total)
                        if page_range == "Specific pages" and page_input:
                            idxs_set = set()
                            for part in page_input.split(","):
                                part = part.strip()
                                if "-" in part:
                                    a, b = part.split("-")
                                    idxs_set.update(range(int(a)-1, min(int(b), total)))
                                elif part:
                                    idxs_set.add(int(part)-1)
                            idxs = sorted(idxs_set)
                        for i in idxs:
                            t = pdf.pages[i].extract_text() or ""
                            text_pages.append(f"--- Page {i+1} ---\n{t}")
                else:
                    reader = PdfReader(io.BytesIO(file.read()))
                    for i, page in enumerate(reader.pages):
                        t = page.extract_text() or ""
                        text_pages.append(f"--- Page {i+1} ---\n{t}")

                full_text = "\n\n".join(text_pages)
                st.text_area("Extracted Text", full_text, height=400)
                word_count = len(full_text.split())
                st.caption(f"{len(text_pages)} pages extracted · {word_count:,} words · {len(full_text):,} characters")
                download_btn(full_text.encode(), "extracted_text.txt", "Download as .txt")
            except Exception as e:
                st.error(f"Error: {e}")
    else:
        st.info("Upload a PDF to extract text from it.")



elif tool == "Extract Tables":
    section_header("Extract Tables", "Pull tables from PDFs and export to CSV or Excel.", premium=True)

    file = st.file_uploader("Upload a PDF containing tables", type="pdf")

    if file:
        import csv
        import pandas as pd

        export_fmt = st.radio("Export format", ["CSV (per table)", "Excel (all tables in one file)"])

        if st.button("Extract Tables", use_container_width=True):
            try:
                all_tables = []
                with pdfplumber.open(io.BytesIO(file.read())) as pdf:
                    for page_num, page in enumerate(pdf.pages):
                        for tbl_idx, tbl in enumerate(page.extract_tables()):
                            if tbl:
                                all_tables.append({"page": page_num + 1, "index": tbl_idx + 1, "data": tbl})

                if not all_tables:
                    st.warning("No tables found in this PDF.")
                else:
                    st.success(f"Found **{len(all_tables)} table(s)**!")
                    for t in all_tables[:3]:
                        st.markdown(f"**Table {t['index']} — Page {t['page']}**")
                        rows = [r for r in t["data"] if r]
                        if rows:
                            df = pd.DataFrame(rows[1:], columns=rows[0]) if len(rows) > 1 else pd.DataFrame(rows)
                            st.dataframe(df, use_container_width=True)

                    if export_fmt.startswith("CSV"):
                        zbuf = io.BytesIO()
                        with zipfile.ZipFile(zbuf, "w") as zf:
                            for t in all_tables:
                                cbuf = io.StringIO()
                                csv.writer(cbuf).writerows(t["data"])
                                zf.writestr(f"page{t['page']}_table{t['index']}.csv", cbuf.getvalue())
                        download_btn(zbuf.getvalue(), "tables_csv.zip", "Download CSVs (ZIP)")
                    else:
                        ebuf = io.BytesIO()
                        with pd.ExcelWriter(ebuf, engine="openpyxl") as writer_xl:
                            for t in all_tables:
                                rows = [r for r in t["data"] if r]
                                if rows:
                                    df = pd.DataFrame(rows[1:], columns=rows[0]) if len(rows) > 1 else pd.DataFrame(rows)
                                    df.to_excel(writer_xl, sheet_name=f"P{t['page']}_T{t['index']}", index=False)
                        download_btn(ebuf.getvalue(), "tables.xlsx", "Download Excel File")
            except Exception as e:
                st.error(f"Error: {e}")
    else:
        st.info("Upload a PDF containing tables to extract them.")



elif tool == "Reorder Pages":
    section_header("Reorder Pages", "Rearrange pages of a PDF in any order you want.")

    file = st.file_uploader("Upload a PDF", type="pdf")

    if file:
        reader = PdfReader(io.BytesIO(file.read()))
        total = len(reader.pages)
        st.info(f"**{file.name}** — {total} pages")

        default_order = ", ".join(str(i+1) for i in range(total))
        new_order = st.text_input(
            f"Enter new page order (1–{total}, comma-separated)",
            value=default_order,
            help="Example: 3,1,2 to move page 3 first. You can also repeat or omit pages."
        )
        st.caption("Tip: Repeat page numbers (e.g. 1,1,2,3) or omit pages to delete them.")

        if st.button("Reorder Pages", use_container_width=True):
            try:
                order = [int(x.strip()) - 1 for x in new_order.split(",") if x.strip()]
                invalid = [i+1 for i in order if i < 0 or i >= total]
                if invalid:
                    st.error(f"Invalid page numbers: {invalid}")
                else:
                    writer = PdfWriter()
                    for i in order:
                        writer.add_page(reader.pages[i])
                    buf = io.BytesIO()
                    writer.write(buf)
                    st.success(f"Pages reordered — new document has {len(order)} pages!")
                    download_btn(buf.getvalue(), "reordered.pdf", "Download Reordered PDF")
            except Exception as e:
                st.error(f"Error: {e}")
    else:
        st.info("Upload a PDF to reorder its pages.")


elif tool == "OCR (Scan to Text)":
    section_header("OCR – Scan to Text", "Extract text from scanned or image-based PDFs using OCR.", premium=True)

    file = st.file_uploader("Upload a scanned PDF", type="pdf")

    if file:
        col1, col2 = st.columns(2)
        language = col1.selectbox("OCR Language", ["eng (English)", "fra (French)", "deu (German)", "spa (Spanish)", "chi_sim (Chinese)", "hin (Hindi)", "ara (Arabic)"])
        dpi = col2.select_slider("Scan resolution", options=[150, 200, 300], value=200)
        lang_code = language.split(" ")[0]

        st.info("OCR may take a while for large documents. Please be patient.")

        if st.button("Run OCR", use_container_width=True):
            try:
                with st.spinner("Converting PDF to images..."):
                    images = convert_from_bytes(file.read(), dpi=dpi)

                all_text = []
                progress = st.progress(0)
                for i, img in enumerate(images):
                    with st.spinner(f"OCR: processing page {i+1}/{len(images)}..."):
                        text = pytesseract.image_to_string(img, lang=lang_code)
                        all_text.append(f"--- Page {i+1} ---\n{text}")
                    progress.progress((i+1) / len(images))

                full_text = "\n\n".join(all_text)
                st.text_area("OCR Result", full_text, height=400)
                st.caption(f"{len(images)} pages · {len(full_text.split()):,} words extracted")
                download_btn(full_text.encode(), "ocr_result.txt", "Download OCR Text")
            except Exception as e:
                st.error(f"Error: {e}. Make sure Tesseract and poppler are installed.")
    else:
        st.info("Upload a scanned PDF to extract text using OCR technology.")


elif tool == "Named Entity Recognition":
    section_header("Named Entity Recognition", "Detect people, organizations, locations, dates and more from your PDF using NLTK.")

    import nltk

    @st.cache_resource(show_spinner="Downloading NLTK models...")
    def download_nltk():
        nltk.download("punkt",                        quiet=True)
        nltk.download("punkt_tab",                    quiet=True)
        nltk.download("averaged_perceptron_tagger",   quiet=True)
        nltk.download("averaged_perceptron_tagger_eng", quiet=True)
        nltk.download("maxent_ne_chunker",            quiet=True)
        nltk.download("maxent_ne_chunker_tab",        quiet=True)
        nltk.download("words",                        quiet=True)

    download_nltk()

    ENTITY_COLORS = {
        "PERSON":       "#4A90D9",
        "ORGANIZATION": "#E67E22",
        "GPE":          "#27AE60",
        "LOCATION":     "#8E44AD",
        "FACILITY":     "#C0392B",
        "GSP":          "#16A085",
        "DATE":         "#F39C12",
        "TIME":         "#2980B9",
    }
    ENTITY_LABELS = {
        "PERSON":       "Person",
        "ORGANIZATION": "Organization",
        "GPE":          "Place (GPE)",
        "LOCATION":     "Location",
        "FACILITY":     "Facility",
        "GSP":          "Geo-Political",
        "DATE":         "Date",
        "TIME":         "Time",
    }

    file = st.file_uploader("Upload a PDF", type="pdf")

    if file:
        col1, col2 = st.columns(2)
        page_limit    = col1.number_input("Max pages to analyse", min_value=1, max_value=100, value=5)
        entity_filter = col2.multiselect(
            "Entity types to show",
            options=list(ENTITY_LABELS.values()),
            default=list(ENTITY_LABELS.values()),
        )

        if st.button("Run NER", use_container_width=True):
            from nltk import word_tokenize, pos_tag, ne_chunk
            from nltk.tree import Tree
            from collections import defaultdict
            import pandas as pd

            with st.spinner("Extracting text from PDF..."):
                text = ""
                with pdfplumber.open(io.BytesIO(file.read())) as pdf:
                    for page in pdf.pages[:int(page_limit)]:
                        t = page.extract_text()
                        if t:
                            text += t + "\n"

            if not text.strip():
                st.warning("No text found in this PDF.")
            else:
                with st.spinner("Running Named Entity Recognition..."):
                    tokens  = word_tokenize(text)
                    tagged  = pos_tag(tokens)
                    chunked = ne_chunk(tagged, binary=False)

                    entities = defaultdict(set)
                    for subtree in chunked:
                        if isinstance(subtree, Tree):
                            etype  = subtree.label()
                            entity = " ".join(word for word, tag in subtree.leaves())
                            if len(entity) > 1:
                                entities[etype].add(entity)

                st.subheader("Summary")
                if entities:
                    cols = st.columns(len(entities))
                    for i, (etype, ents) in enumerate(entities.items()):
                        cols[i].metric(ENTITY_LABELS.get(etype, etype), len(ents))
                else:
                    st.info("No named entities found.")

                st.divider()
                st.subheader("Entities Found")

                rev_labels     = {v: k for k, v in ENTITY_LABELS.items()}
                selected_types = {rev_labels.get(lbl, lbl) for lbl in entity_filter}

                rows = []
                for etype, ents in sorted(entities.items()):
                    if etype not in selected_types:
                        continue
                    label = ENTITY_LABELS.get(etype, etype)
                    st.markdown(f"**{label}**")
                    tags_html = " ".join(
                        f'<span style="background:{ENTITY_COLORS.get(etype,"#888")};'
                        f'color:white;padding:3px 10px;border-radius:12px;'
                        f'margin:3px;display:inline-block;font-size:0.88rem">{e}</span>'
                        for e in sorted(ents)
                    )
                    st.markdown(tags_html, unsafe_allow_html=True)
                    st.write("")
                    for e in sorted(ents):
                        rows.append({"Entity": e, "Type": label})

                if rows:
                    st.divider()
                    st.subheader("Export")
                    df = pd.DataFrame(rows).sort_values(["Type", "Entity"]).reset_index(drop=True)
                    st.dataframe(df, use_container_width=True)
                    download_btn(df.to_csv(index=False).encode(), "named_entities.csv", "Download as CSV")
    else:
        st.info("Upload a PDF to extract named entities from it.")